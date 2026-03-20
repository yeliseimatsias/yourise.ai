import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


class ReportGenerator:
    def __init__(self, session_info: dict, changes: list):
        """
        :param session_info: Метаданные сессии (названия файлов и т.д.)
        :param changes: Список изменений, где объединены данные UI (тексты) и ответ от LLM
        """
        self.session = session_info
        self.changes = changes
        self.generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Автоматически считаем статистику по новым данным
        self.stats = {"total_changes": len(changes), "red": 0, "yellow": 0, "green": 0}
        for ch in changes:
            risk = ch.get("overall_risk", "green")
            if risk in self.stats:
                self.stats[risk] += 1

    def _get_short_advice(self, risk_level: str) -> str:
        """Генерирует статус в зависимости от уровня риска."""
        advices = {
            "red": "Критическая ошибка. Требуется исправление.",
            "yellow": "Внимание. Потенциальный риск, требуется проверка юристом.",
            "green": "Допустимо. Соответствует законодательству."
        }
        return advices.get(risk_level, "Требуется ручная проверка.")

    def generate_json_report(self) -> dict:
        """Формирует структурированный JSON для UI."""
        session_id = self.session.get("id", "unknown_session")

        report = {
            "session_id": session_id,
            "generated_at": self.generated_at,
            "summary": {
                "documents": {
                    "old": self.session.get("old_document", "Неизвестно"),
                    "new": self.session.get("new_document", "Неизвестно"),
                },
                "stats": self.stats,
            },
            "changes_by_risk": {"red": [], "yellow": [], "green": []},
            "download_links": {
                "json": f"/api/reports/{session_id}/json",
                "docx": f"/api/reports/{session_id}/docx"
            }
        }

        for change in self.changes:
            risk = change.get("overall_risk", "green")

            # Берем первую валидацию из ответа LLM (если она есть)
            validations = change.get("validation_results", [])
            validation = validations[0] if validations else {}

            # 🔥 ДОБАВЛЯЕМ old_number
            entry = {
                "change_id": change.get("change_id"),
                "number": change.get("element_number", "—"),           # новый номер
                "old_number": change.get("old_number"),                # старый номер (может быть None)
                "type": change.get("change_type", "modified"),
                "old_text": change.get("old_text", "—"),
                "new_text": change.get("new_text", "—"),
                "issue": {
                    "explanation": validation.get("explanation",
                                                  change.get("overall_explanation", "Плановое изменение")),
                    "law_reference": validation.get("law_reference", "Не требуется"),
                    "suggestion": validation.get("suggestion", ""),
                    "advice": self._get_short_advice(risk)
                }
            }
            report["changes_by_risk"][risk].append(entry)

        return report

    def _add_risk_section(self, doc, title, changes, color_rgb):
        """Вспомогательный метод для отрисовки секции риска в DOCX."""
        if not changes:
            return

        heading = doc.add_heading(title, level=1)
        run = heading.runs[0]
        run.font.color.rgb = color_rgb

        for item in changes:
            # 🔥 Добавляем отображение старого номера в заголовок
            if item.get('old_number') and item['old_number'] != item['number']:
                display_number = f"{item['old_number']} → {item['number']}"
            else:
                display_number = item['number']

            doc.add_heading(f"Пункт {display_number} ({item['type']})", level=2)

            # Таблица сравнения текста
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            t_cells = table.rows[0].cells
            t_cells[0].text = f"БЫЛО:\n{item['old_text']}"
            t_cells[1].text = f"СТАЛО:\n{item['new_text']}"

            # Аналитика от LLM
            p_analys = doc.add_paragraph()
            p_analys.add_run("\nАнализ LLM: ").bold = True
            p_analys.add_run(item['issue']['explanation'])

            p_law = doc.add_paragraph()
            p_law.add_run("Основание: ").bold = True
            p_law.add_run(item['issue']['law_reference'])

            # Добавляем рекомендацию из LLM
            if item['issue']['suggestion']:
                p_sug = doc.add_paragraph()
                p_sug.add_run("Решение: ").bold = True
                p_sug.add_run(item['issue']['suggestion'])

            # Краткий статус (выделен цветом и курсивом)
            p_advice = doc.add_paragraph()
            run_adv = p_advice.add_run("Статус: ")
            run_adv.bold = True
            run_adv.font.color.rgb = color_rgb
            p_advice.add_run(item['issue']['advice']).italic = True

            doc.add_paragraph()  # Отступ между пунктами

    def generate_docx_report(self) -> BytesIO:
        """Генерирует полный DOCX файл с разделением по зонам."""
        json_data = self.generate_json_report()
        doc = Document()

        # Титульная часть
        title = doc.add_heading('Юридическое заключение по результатам сравнения', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Дата отчета: {json_data['generated_at']}")
        doc.add_paragraph(f"Документ: {json_data['summary']['documents']['new']}")
        doc.add_paragraph("-" * 30)

        # 1. КРАСНАЯ ЗОНА
        self._add_risk_section(
            doc, "1. КРИТИЧЕСКИЕ РИСКИ (КРАСНАЯ ЗОНА)",
            json_data['changes_by_risk']['red'], RGBColor(200, 0, 0)
        )

        # 2. ЖЕЛТАЯ ЗОНА
        self._add_risk_section(
            doc, "2. СРЕДНИЕ РИСКИ (ЖЕЛТАЯ ЗОНА)",
            json_data['changes_by_risk']['yellow'], RGBColor(218, 165, 32)
        )

        # 3. ЗЕЛЕНАЯ ЗОНА
        self._add_risk_section(
            doc, "3. НИЗКИЙ РИСК / ТЕХНИЧЕСКИЕ ПРАВКИ",
            json_data['changes_by_risk']['green'], RGBColor(34, 139, 34)
        )

        target_stream = BytesIO()
        doc.save(target_stream)
        target_stream.seek(0)
        return target_stream